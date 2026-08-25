"""
Tests for the format-agnostic ingestion architecture.

Covers:
- Format registry
- Format detection
- Document adapters
- Validation
- Unsupported format rejection
- Security validation
"""

from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path
from unittest.mock import patch

from PIL import Image

from django.test import TestCase, override_settings

from ocr.adapters import (
    CSVAdapter,
    DOCXAdapter,
    ImageAdapter,
    PDFAdapter,
    SpreadsheetAdapter,
    TextAdapter,
    get_adapter,
)
from ocr.exceptions import (
    DocumentProcessingException,
    UnsupportedFormatException,
)
from ocr.formats import (
    SUPPORTED_FORMATS,
    FormatCategory,
    detect_format,
    get_supported_extensions,
    get_supported_mime_types,
    is_supported_extension,
)
from ocr.validators import (
    ALLOWED_EXTENSIONS,
    ALLOWED_MIME_TYPES,
    MAX_FILE_SIZE,
    validate_extension,
    validate_file_format,
    validate_file_size,
    validate_mime_type,
)


# ==================================================================
# Format Registry Tests
# ==================================================================

class FormatRegistryTests(TestCase):
    """Tests for ocr.formats."""

    def test_supported_extensions_returns_sorted_list(self):
        extensions = get_supported_extensions()
        self.assertEqual(extensions, sorted(extensions))
        self.assertIn('pdf', extensions)
        self.assertIn('png', extensions)
        self.assertIn('xlsx', extensions)
        self.assertIn('csv', extensions)
        self.assertIn('txt', extensions)

    def test_supported_mime_types_returns_set(self):
        mimes = get_supported_mime_types()
        self.assertIsInstance(mimes, frozenset)
        self.assertIn('application/pdf', mimes)
        self.assertIn('image/png', mimes)

    def test_is_supported_extension(self):
        self.assertTrue(is_supported_extension('pdf'))
        self.assertTrue(is_supported_extension('PDF'))
        self.assertTrue(is_supported_extension('xlsx'))
        self.assertTrue(is_supported_extension('docx'))
        self.assertFalse(is_supported_extension('doc'))
        self.assertFalse(is_supported_extension('exe'))

    def test_allowed_extensions_alias_matches_registry(self):
        self.assertEqual(set(ALLOWED_EXTENSIONS), set(SUPPORTED_FORMATS.keys()))

    def test_allowed_mime_types_alias_matches_registry(self):
        expected = get_supported_mime_types()
        self.assertEqual(ALLOWED_MIME_TYPES, expected)

    def test_registry_has_no_audio_or_video(self):
        for entry in SUPPORTED_FORMATS.values():
            mime = entry.mime_types[0]
            self.assertFalse(mime.startswith('audio/'))
            self.assertFalse(mime.startswith('video/'))

    def test_registry_has_no_executable_or_archive(self):
        unsafe = {
            'application/x-msdownload',
            'application/x-executable',
            'application/x-sharedlib',
            'application/x-compress',
            'application/x-gzip',
            'application/zip',
            'application/x-rar-compressed',
        }
        for entry in SUPPORTED_FORMATS.values():
            for mime in entry.mime_types:
                self.assertNotIn(mime, unsafe)


# ==================================================================
# Format Detection Tests
# ==================================================================

class FormatDetectionTests(TestCase):
    """Tests for ocr.formats.detect_format."""

    def test_detect_pdf(self):
        entry = detect_format('invoice.pdf')
        self.assertEqual(entry.extension, 'pdf')
        self.assertEqual(entry.category, FormatCategory.DOCUMENT)

    def test_detect_pdf_uppercase(self):
        entry = detect_format('invoice.PDF')
        self.assertEqual(entry.extension, 'pdf')

    def test_detect_image_jpg(self):
        entry = detect_format('photo.jpg')
        self.assertEqual(entry.extension, 'jpg')
        self.assertEqual(entry.category, FormatCategory.IMAGE)

    def test_detect_image_png(self):
        entry = detect_format('photo.png')
        self.assertEqual(entry.extension, 'png')
        self.assertEqual(entry.category, FormatCategory.IMAGE)

    def test_detect_image_webp(self):
        entry = detect_format('photo.webp')
        self.assertEqual(entry.extension, 'webp')

    def test_detect_image_gif(self):
        entry = detect_format('photo.gif')
        self.assertEqual(entry.extension, 'gif')

    def test_detect_image_bmp(self):
        entry = detect_format('photo.bmp')
        self.assertEqual(entry.extension, 'bmp')

    def test_detect_image_tif(self):
        entry = detect_format('photo.tif')
        self.assertEqual(entry.extension, 'tif')

    def test_detect_image_tiff(self):
        entry = detect_format('photo.tiff')
        self.assertEqual(entry.extension, 'tiff')

    def test_detect_xlsx(self):
        entry = detect_format('data.xlsx')
        self.assertEqual(entry.extension, 'xlsx')
        self.assertEqual(entry.category, FormatCategory.SPREADSHEET)

    def test_detect_csv(self):
        entry = detect_format('data.csv')
        self.assertEqual(entry.extension, 'csv')
        self.assertEqual(entry.category, FormatCategory.SPREADSHEET)

    def test_detect_txt(self):
        entry = detect_format('notes.txt')
        self.assertEqual(entry.extension, 'txt')
        self.assertEqual(entry.category, FormatCategory.TEXT)

    def test_detect_docx(self):
        entry = detect_format('invoice.docx')
        self.assertEqual(entry.extension, 'docx')
        self.assertEqual(entry.category, FormatCategory.DOCUMENT)

    def test_detect_unsupported_raises(self):
        with self.assertRaises(UnsupportedFormatException):
            detect_format('invoice.doc')

    def test_detect_executable_raises(self):
        with self.assertRaises(UnsupportedFormatException):
            detect_format('malware.exe')

    def test_detect_no_extension_raises(self):
        with self.assertRaises(UnsupportedFormatException):
            detect_format('noextension')

    def test_detect_mime_mismatch_raises(self):
        with self.assertRaises(UnsupportedFormatException):
            detect_format('invoice.pdf.exe')


# ==================================================================
# Validator Tests
# ==================================================================

class ValidatorTests(TestCase):
    """Tests for ocr.validators."""

    def test_max_file_size_constant(self):
        self.assertEqual(MAX_FILE_SIZE, 10 * 1024 * 1024)

    def test_validate_file_size_valid(self):
        validate_file_size(0)
        validate_file_size(MAX_FILE_SIZE)

    def test_validate_file_size_over_limit(self):
        with self.assertRaises(Exception):
            validate_file_size(MAX_FILE_SIZE + 1)

    def test_validate_extension_supported(self):
        for ext in ['pdf', 'png', 'jpg', 'jpeg', 'xlsx', 'csv', 'txt']:
            validate_extension(f'invoice.{ext}')

    def test_validate_extension_unsupported(self):
        with self.assertRaises(Exception):
            validate_extension('invoice.doc')

    def test_validate_mime_type_supported(self):
        for mime in ALLOWED_MIME_TYPES:
            validate_mime_type(mime)

    def test_validate_mime_type_unsupported(self):
        with self.assertRaises(Exception):
            validate_mime_type('application/msword')

    def test_validate_file_format_valid_pdf(self):
        pdf_path = _create_temp_pdf()
        try:
            result = validate_file_format(pdf_path, 'invoice.pdf')
            self.assertTrue(result['is_valid'])
        finally:
            os.unlink(pdf_path)

    def test_validate_file_format_valid_image(self):
        img_path = _create_temp_png()
        try:
            result = validate_file_format(img_path, 'photo.png')
            self.assertTrue(result['is_valid'])
        finally:
            os.unlink(img_path)

    def test_validate_file_format_invalid_extension(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp:
            tmp.write(b'doc content')
            tmp_path = tmp.name
        try:
            result = validate_file_format(tmp_path, 'invoice.doc')
            self.assertFalse(result['is_valid'])
            self.assertTrue(
                any('not allowed' in e or 'not supported' in e for e in result['errors'])
            )
        finally:
            os.unlink(tmp_path)

    def test_validate_file_format_unsupported_format(self):
        result = validate_file_format('/tmp/test.exe', 'malware.exe')
        self.assertFalse(result['is_valid'])

    def test_validate_file_format_oversized(self):
        path = _create_temp_pdf()
        try:
            with self.assertRaises(Exception):
                validate_file_format(path, 'invoice.pdf', max_size=1)
        finally:
            os.unlink(path)


# ==================================================================
# Adapter Tests
# ==================================================================

class PDFAdapterTests(TestCase):
    """Tests for ocr.adapters.PDFAdapter."""

    def test_normalize_returns_images(self):
        pdf_path = _create_temp_pdf()
        try:
            adapter = PDFAdapter(file_path=pdf_path, upload_id='test-pdf')
            result = adapter.normalize()
            self.assertIn('pages', result)
            self.assertIn('metadata', result)
            self.assertEqual(result['format'], 'pdf')
            self.assertGreater(result['metadata']['page_count'], 0)
        finally:
            os.unlink(pdf_path)
            from ocr.pdf_processor import pdf_processor
            pdf_processor.cleanup('test-pdf')

    def test_normalize_invalid_pdf_raises(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            tmp.write(b'not a pdf')
            tmp_path = tmp.name
        try:
            adapter = PDFAdapter(file_path=tmp_path, upload_id='test-invalid')
            with self.assertRaises(DocumentProcessingException):
                adapter.normalize()
        finally:
            try:
                os.unlink(tmp_path)
            except PermissionError:
                pass  # Windows may lock the file briefly after fitz error


class ImageAdapterTests(TestCase):
    """Tests for ocr.adapters.ImageAdapter."""

    def test_normalize_valid_image(self):
        img_path = _create_temp_png()
        try:
            adapter = ImageAdapter(file_path=img_path, upload_id='test-img')
            result = adapter.normalize()
            self.assertIn('pages', result)
            self.assertIn('metadata', result)
            self.assertEqual(result['format'], 'png')
            self.assertEqual(len(result['pages']), 1)
        finally:
            os.unlink(img_path)

    def test_normalize_invalid_image_raises(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
            tmp.write(b'not an image')
            tmp_path = tmp.name
        try:
            adapter = ImageAdapter(file_path=tmp_path, upload_id='test-invalid')
            with self.assertRaises(DocumentProcessingException):
                adapter.normalize()
        finally:
            os.unlink(tmp_path)

    def test_validate_security_valid_image(self):
        img_path = _create_temp_png()
        try:
            adapter = ImageAdapter(file_path=img_path, upload_id='test-img')
            adapter.validate_security()  # Should not raise
        finally:
            os.unlink(img_path)

    def test_validate_security_invalid_image_raises(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
            tmp.write(b'not an image')
            tmp_path = tmp.name
        try:
            adapter = ImageAdapter(file_path=tmp_path, upload_id='test-invalid')
            with self.assertRaises(DocumentProcessingException):
                adapter.validate_security()
        finally:
            os.unlink(tmp_path)


class SpreadsheetAdapterTests(TestCase):
    """Tests for ocr.adapters.SpreadsheetAdapter."""

    def test_normalize_xlsx(self):
        xlsx_path = _create_temp_xlsx()
        try:
            adapter = SpreadsheetAdapter(file_path=xlsx_path, upload_id='test-xlsx')
            result = adapter.normalize()
            self.assertIn('text', result)
            self.assertIn('metadata', result)
            self.assertEqual(result['format'], 'xlsx')
            self.assertIn('Sheet1', result['text'])
        finally:
            os.unlink(xlsx_path)

    def test_normalize_xlsx_has_headers_and_rows(self):
        xlsx_path = _create_temp_xlsx()
        try:
            adapter = SpreadsheetAdapter(file_path=xlsx_path, upload_id='test-xlsx')
            result = adapter.normalize()
            self.assertIn('Invoice No', result['text'])
            self.assertIn('INV-001', result['text'])
        finally:
            os.unlink(xlsx_path)


class CSVAdapterTests(TestCase):
    """Tests for ocr.adapters.CSVAdapter."""

    def test_normalize_csv(self):
        csv_path = _create_temp_csv()
        try:
            adapter = CSVAdapter(file_path=csv_path, upload_id='test-csv')
            result = adapter.normalize()
            self.assertIn('text', result)
            self.assertIn('metadata', result)
            self.assertEqual(result['format'], 'csv')
            self.assertIn('Invoice No', result['text'])
        finally:
            os.unlink(csv_path)

    def test_normalize_csv_has_headers_and_rows(self):
        csv_path = _create_temp_csv()
        try:
            adapter = CSVAdapter(file_path=csv_path, upload_id='test-csv')
            result = adapter.normalize()
            self.assertIn('INV-001', result['text'])
            self.assertEqual(result['metadata']['row_count'], 1)
        finally:
            os.unlink(csv_path)

    def test_validate_security_rejects_zip(self):
        import zipfile
        with tempfile.NamedTemporaryFile(delete=False, suffix='.csv') as tmp:
            with zipfile.ZipFile(tmp.name, 'w') as zf:
                zf.writestr('test.txt', 'hello')
            tmp_path = tmp.name
        try:
            adapter = CSVAdapter(file_path=tmp_path, upload_id='test-csv')
            with self.assertRaises(DocumentProcessingException):
                adapter.validate_security()
        finally:
            os.unlink(tmp_path)


class TextAdapterTests(TestCase):
    """Tests for ocr.adapters.TextAdapter."""

    def test_normalize_txt(self):
        txt_path = _create_temp_txt()
        try:
            adapter = TextAdapter(file_path=txt_path, upload_id='test-txt')
            result = adapter.normalize()
            self.assertIn('text', result)
            self.assertIn('metadata', result)
            self.assertEqual(result['format'], 'txt')
            self.assertIn('Hello World', result['text'])
        finally:
            os.unlink(txt_path)

    def test_normalize_rejects_binary(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.txt') as tmp:
            tmp.write(b'Hello\x00World')
            tmp_path = tmp.name
        try:
            adapter = TextAdapter(file_path=tmp_path, upload_id='test-txt')
            with self.assertRaises(DocumentProcessingException):
                adapter.normalize()
        finally:
            os.unlink(tmp_path)

    def test_normalize_rejects_oversized(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.txt') as tmp:
            tmp.write(b'x' * (5 * 1024 * 1024 + 1))
            tmp_path = tmp.name
        try:
            adapter = TextAdapter(file_path=tmp_path, upload_id='test-txt')
            with self.assertRaises(DocumentProcessingException):
                adapter.normalize()
        finally:
            os.unlink(tmp_path)


class DOCXAdapterTests(TestCase):
    """Tests for ocr.adapters.DOCXAdapter."""

    def test_normalize_docx(self):
        docx_path = _create_temp_docx()
        try:
            adapter = DOCXAdapter(file_path=docx_path, upload_id='test-docx')
            result = adapter.normalize()
            self.assertIn('text', result)
            self.assertIn('metadata', result)
            self.assertEqual(result['format'], 'docx')
            self.assertIn('Hello World', result['text'])
        finally:
            os.unlink(docx_path)

    def test_normalize_docx_has_paragraphs_and_tables(self):
        docx_path = _create_temp_docx_with_table()
        try:
            adapter = DOCXAdapter(file_path=docx_path, upload_id='test-docx-table')
            result = adapter.normalize()
            self.assertIn('Hello World', result['text'])
            self.assertIn('TABLE', result['text'])
            self.assertIn('Cell A', result['text'])
            self.assertEqual(result['metadata']['paragraph_count'], 1)
            self.assertEqual(result['metadata']['table_count'], 1)
        finally:
            os.unlink(docx_path)

    def test_normalize_docx_returns_image_pages(self):
        docx_path = _create_temp_docx()
        try:
            adapter = DOCXAdapter(file_path=docx_path, upload_id='test-docx-img')
            result = adapter.normalize()
            self.assertIn('pages', result)
            self.assertTrue(len(result['pages']) > 0)
        finally:
            os.unlink(docx_path)

    def test_validate_security_valid_docx(self):
        docx_path = _create_temp_docx()
        try:
            adapter = DOCXAdapter(file_path=docx_path, upload_id='test-docx')
            adapter.validate_security()  # Should not raise
        finally:
            os.unlink(docx_path)

    def test_validate_security_invalid_docx_raises(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(b'not a docx file')
            tmp_path = tmp.name
        try:
            adapter = DOCXAdapter(file_path=tmp_path, upload_id='test-invalid')
            with self.assertRaises(DocumentProcessingException):
                adapter.validate_security()
        finally:
            os.unlink(tmp_path)

    def test_normalize_docx_embedded_image_format(self):
        docx_path = _create_temp_docx_with_embedded_image()
        try:
            adapter = DOCXAdapter(file_path=docx_path, upload_id='test-docx-img')
            result = adapter.normalize()
            self.assertIn('pages', result)
            self.assertTrue(len(result['pages']) > 1)
            for page in result['pages'][1:]:
                self.assertTrue(page.endswith('.jpg') or page.endswith('.png'))
        finally:
            os.unlink(docx_path)

    def test_normalize_docx_preserves_order(self):
        docx_path = _create_temp_docx_with_order()
        try:
            adapter = DOCXAdapter(file_path=docx_path, upload_id='test-docx-order')
            result = adapter.normalize()
            text = result['text']
            para1_idx = text.index('Paragraph 1')
            table_idx = text.index('[TABLE]')
            para2_idx = text.index('Paragraph 2')
            self.assertLess(para1_idx, table_idx)
            self.assertLess(table_idx, para2_idx)
        finally:
            os.unlink(docx_path)


class AdapterFactoryTests(TestCase):
    """Tests for ocr.adapters.get_adapter."""

    def test_get_adapter_pdf(self):
        pdf_path = _create_temp_pdf()
        try:
            adapter = get_adapter(pdf_path, 'test')
            self.assertIsInstance(adapter, PDFAdapter)
        finally:
            os.unlink(pdf_path)

    def test_get_adapter_docx(self):
        docx_path = _create_temp_docx()
        try:
            adapter = get_adapter(docx_path, 'test')
            self.assertIsInstance(adapter, DOCXAdapter)
        finally:
            os.unlink(docx_path)

    def test_get_adapter_png(self):
        img_path = _create_temp_png()
        try:
            adapter = get_adapter(img_path, 'test')
            self.assertIsInstance(adapter, ImageAdapter)
        finally:
            os.unlink(img_path)

    def test_get_adapter_xlsx(self):
        xlsx_path = _create_temp_xlsx()
        try:
            adapter = get_adapter(xlsx_path, 'test')
            self.assertIsInstance(adapter, SpreadsheetAdapter)
        finally:
            os.unlink(xlsx_path)

    def test_get_adapter_csv(self):
        csv_path = _create_temp_csv()
        try:
            adapter = get_adapter(csv_path, 'test')
            self.assertIsInstance(adapter, CSVAdapter)
        finally:
            os.unlink(csv_path)

    def test_get_adapter_txt(self):
        txt_path = _create_temp_txt()
        try:
            adapter = get_adapter(txt_path, 'test')
            self.assertIsInstance(adapter, TextAdapter)
        finally:
            os.unlink(txt_path)

    def test_get_adapter_unsupported_raises(self):
        with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp:
            tmp.write(b'doc content')
            tmp_path = tmp.name
        try:
            with self.assertRaises(UnsupportedFormatException):
                get_adapter(tmp_path, 'test')
        finally:
            os.unlink(tmp_path)


# ==================================================================
# Helper functions
# ==================================================================

def _create_temp_pdf() -> str:
    """Create a minimal valid PDF file and return its path."""
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), 'Test Invoice')
    pdf_bytes = doc.tobytes()
    doc.close()
    fd, path = tempfile.mkstemp(suffix='.pdf')
    with os.fdopen(fd, 'wb') as f:
        f.write(pdf_bytes)
    return path


def _create_temp_png() -> str:
    """Create a valid PNG file and return its path."""
    img = Image.new('RGB', (100, 100), color='red')
    fd, path = tempfile.mkstemp(suffix='.png')
    img.save(path, format='PNG')
    os.close(fd)
    return path


def _create_temp_xlsx() -> str:
    """Create a minimal XLSX file and return its path."""
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Sheet1'
    ws.append(['Invoice No', 'Date', 'Vendor', 'Total'])
    ws.append(['INV-001', '2024-01-01', 'Acme Corp', '100.00'])
    fd, path = tempfile.mkstemp(suffix='.xlsx')
    wb.save(path)
    wb.close()
    os.close(fd)
    return path


def _create_temp_csv() -> str:
    """Create a CSV file and return its path."""
    fd, path = tempfile.mkstemp(suffix='.csv')
    with os.fdopen(fd, 'w', encoding='utf-8') as f:
        f.write('Invoice No,Date,Vendor,Total\n')
        f.write('INV-001,2024-01-01,Acme Corp,100.00\n')
    return path


def _create_temp_txt() -> str:
    """Create a text file and return its path."""
    fd, path = tempfile.mkstemp(suffix='.txt')
    with os.fdopen(fd, 'w', encoding='utf-8') as f:
        f.write('Hello World\n')
    return path


def _create_temp_docx() -> str:
    """Create a minimal DOCX file and return its path."""
    from docx import Document
    doc = Document()
    doc.add_paragraph('Hello World')
    doc.add_paragraph('Test Invoice')
    fd, path = tempfile.mkstemp(suffix='.docx')
    doc.save(path)
    os.close(fd)
    return path


def _create_temp_docx_with_table() -> str:
    """Create a DOCX file with a table and return its path."""
    from docx import Document
    doc = Document()
    doc.add_paragraph('Hello World')
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = 'Cell A'
    table.cell(0, 1).text = 'Cell B'
    fd, path = tempfile.mkstemp(suffix='.docx')
    doc.save(path)
    os.close(fd)
    return path


def _create_temp_docx_with_embedded_image() -> str:
    """Create a DOCX file with an embedded JPEG image."""
    from docx import Document
    from docx.shared import Inches
    doc = Document()
    doc.add_paragraph('Invoice with image')
    # Add a simple red square image
    img_path = _create_temp_png()
    doc.add_picture(img_path, width=Inches(1))
    fd, path = tempfile.mkstemp(suffix='.docx')
    doc.save(path)
    os.close(fd)
    os.unlink(img_path)
    return path


def _create_temp_docx_with_order() -> str:
    """Create a DOCX file with interleaved paragraphs and tables."""
    from docx import Document
    doc = Document()
    doc.add_paragraph('Paragraph 1')
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = 'Table Cell 1'
    table.cell(0, 1).text = 'Table Cell 2'
    doc.add_paragraph('Paragraph 2')
    fd, path = tempfile.mkstemp(suffix='.docx')
    doc.save(path)
    os.close(fd)
    return path
