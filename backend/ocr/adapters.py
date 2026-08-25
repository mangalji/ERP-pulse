"""
Document adapters for format-agnostic OCR ingestion.

Each adapter converts a specific source format into a normalized
representation that the downstream extraction pipeline can consume.
Adapters are selected via the format registry (``ocr.formats``).

Design goals:
- Keep format-specific parsing isolated from the extraction contract.
- Never execute macros, scripts, or embedded code.
- Preserve enough structure for the AI extraction layer to work.
- Fail loudly with clear errors on unsupported or corrupt inputs.
"""

from __future__ import annotations

import csv
import os
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

from ocr.exceptions import (
    DocumentProcessingException,
    UnsupportedFormatException,
)
from ocr.formats import FormatEntry
from ocr.utils import logger


# ------------------------------------------------------------------
# Base adapter
# ------------------------------------------------------------------

class BaseDocumentAdapter(ABC):
    """
    Abstract base for all document adapters.
    """

    format_entry: FormatEntry = None

    def __init__(self, file_path: str | Path, upload_id: str, format_entry: FormatEntry = None):
        self.file_path = Path(file_path)
        self.upload_id = upload_id
        if format_entry is not None:
            self.format_entry = format_entry
        elif self.format_entry is None:
            from ocr.formats import detect_format
            try:
                self.format_entry = detect_format(file_path)
            except Exception:
                pass

    @abstractmethod
    def normalize(self) -> dict[str, Any]:
        """
        Convert the source document into a normalized representation.

        Returns:
            A dict with at least:
            - ``text``: the primary text content for AI extraction
            - ``pages``: list of image paths for the extraction pipeline
            - ``metadata``: format-specific metadata
            - ``format``: the source format identifier
        """

    def cleanup(self) -> None:
        """
        Remove any temporary files created during normalization.
        """

    def validate_security(self) -> None:
        """
        Perform security checks specific to this format.

        Raises on any indicator of malicious content.
        """


# ------------------------------------------------------------------
# PDF adapter
# ------------------------------------------------------------------

class PDFAdapter(BaseDocumentAdapter):
    """
    Convert PDF files to images using the existing pdf_processor.

    Preserves existing PDF behavior: multi-page support, DPI control,
    temporary directory isolation, and cleanup.
    """

    def normalize(self) -> dict[str, Any]:
        from ocr.exceptions import PDFProcessingException
        from ocr.pdf_processor import pdf_processor

        try:
            page_count = pdf_processor.get_page_count(self.file_path)
            images = pdf_processor.convert_to_images(self.file_path, self.upload_id)
        except PDFProcessingException as exc:
            raise DocumentProcessingException(
                f'Failed to process PDF: {exc}'
            ) from exc

        return {
            'text': '',
            'pages': [str(p) for p in images],
            'metadata': {
                'page_count': page_count,
                'format': self.format_entry.extension if self.format_entry else 'pdf',
            },
            'format': self.format_entry.extension if self.format_entry else 'pdf',
        }


# ------------------------------------------------------------------
# DOCX adapter
# ------------------------------------------------------------------

class DOCXAdapter(BaseDocumentAdapter):
    """
    Convert DOCX documents to a structured text representation.

    Extracts paragraphs, tables, headers, footers, and embedded images.
    Renders the content to an image for the vision pipeline.
    """

    def normalize(self) -> dict[str, Any]:
        try:
            import docx
            from docx.oxml.ns import qn
        except ImportError as exc:
            raise UnsupportedFormatException(
                'python-docx is not installed; DOCX processing is unavailable.'
            ) from exc

        doc = docx.Document(str(self.file_path))
        text_parts = []

        # Extract headers and footers first
        for section in doc.sections:
            header = section.header
            footer = section.footer
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(f'[HEADER] {paragraph.text}')
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(f'[FOOTER] {paragraph.text}')

        # Extract body in document order (paragraphs and tables interleaved)
        for element in doc.element.body:
            if element.tag == qn('w:p'):
                para = docx.text.paragraph.Paragraph(element, doc)
                if para.text.strip():
                    text_parts.append(para.text)
            elif element.tag == qn('w:tbl'):
                table = docx.table.Table(element, doc)
                text_parts.append('[TABLE]')
                for row in table.rows:
                    row_data = ' | '.join(cell.text.strip() for cell in row.cells)
                    text_parts.append(row_data)

        # Extract embedded images
        image_paths = []
        for rel in doc.part.rels.values():
            if 'image' in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    ext = self._detect_image_extension(image_data)
                    if ext is None:
                        logger.warning(
                            'Skipping DOCX embedded image with unknown format'
                        )
                        continue
                    
                    image_dir = Path('/tmp/ocr_docx_images')
                    image_dir.mkdir(parents=True, exist_ok=True)
                    image_path = image_dir / f'{self.upload_id}_{len(image_paths)}{ext}'
                    with open(image_path, 'wb') as f:
                        f.write(image_data)
                    image_paths.append(str(image_path))
                except Exception as exc:
                    logger.warning('Failed to extract DOCX embedded image: %s', exc)

        text = '\n'.join(text_parts)
        rendered_image = self._render_text_to_image(text)

        pages = [rendered_image] if rendered_image else []
        pages.extend(image_paths)

        return {
            'text': text,
            'pages': pages,
            'metadata': {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'image_count': len(image_paths),
                'format': 'docx',
            },
            'format': 'docx',
        }

    @staticmethod
    def _detect_image_extension(data: bytes) -> str | None:
        """
        Detect image format from magic bytes.
        """
        if data.startswith(b'\xFF\xD8\xFF'):
            return '.jpg'
        if data.startswith(b'\x89PNG\r\n\x1a\n'):
            return '.png'
        if data.startswith(b'GIF87a') or data.startswith(b'GIF89a'):
            return '.gif'
        if data.startswith(b'BM'):
            return '.bmp'
        if data.startswith(b'\x49\x49\x2A\x00') or data.startswith(b'\x4D\x4D\x00\x2A'):
            return '.tiff'
        if data.startswith(b'RIFF') and data[8:12] == b'WEBP':
            return '.webp'
        return None

    def _render_text_to_image(self, text: str) -> str | None:
        """
        Render extracted text to a PNG image for the vision pipeline.
        """
        try:
            from PIL import Image, ImageDraw, ImageFont

            font_size = 14
            margin = 20
            line_height = font_size + 4

            lines = text.split('\n')
            max_width = max(len(line) for line in lines) if lines else 80
            width = max(800, max_width * (font_size // 2) + margin * 2)
            height = max(600, len(lines) * line_height + margin * 2)

            img = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(img)

            try:
                font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', font_size)
            except Exception:
                font = ImageFont.load_default()

            y = margin
            for line in lines:
                draw.text((margin, y), line, fill='black', font=font)
                y += line_height

            output_dir = Path('/tmp/ocr_rendered')
            output_dir.mkdir(parents=True, exist_ok=True)
            image_path = output_dir / f'{self.upload_id}_docx.png'
            img.save(str(image_path))
            return str(image_path)
        except Exception as exc:
            logger.warning('Failed to render DOCX text to image: %s', exc)
            return None

    def validate_security(self) -> None:
        try:
            import docx
            doc = docx.Document(str(self.file_path))
            doc._part  # Access internal part to trigger XML parsing
        except Exception as exc:
            raise DocumentProcessingException(
                f'DOCX security validation failed: {exc}'
            ) from exc

    def cleanup(self) -> None:
        """
        Remove any temporary files created during normalization.
        """
        try:
            image_dir = Path('/tmp/ocr_docx_images')
            if image_dir.exists():
                for f in image_dir.glob(f'{self.upload_id}_*'):
                    f.unlink()
        except Exception:
            pass


# ------------------------------------------------------------------
# Image adapter
# ------------------------------------------------------------------

class ImageAdapter(BaseDocumentAdapter):
    """
    Validate and pass through image files.

    Uses Pillow for content validation (size, dimensions, format).
    Does not trust the file extension alone.
    """

    def normalize(self) -> dict[str, Any]:
        from PIL import Image
        from PIL import UnidentifiedImageError

        try:
            with Image.open(self.file_path) as img:
                img.verify()
        except UnidentifiedImageError as exc:
            raise DocumentProcessingException(
                f'File is not a valid image: {self.file_path}'
            ) from exc
        except Exception as exc:
            raise DocumentProcessingException(
                f'Failed to validate image: {exc}'
            ) from exc

        with Image.open(self.file_path) as img:
            width, height = img.size
            if width == 0 or height == 0:
                raise DocumentProcessingException(
                    f'Image has zero dimensions: {self.file_path}'
                )

        return {
            'text': '',
            'pages': [str(self.file_path)],
            'metadata': {
                'width': width,
                'height': height,
                'format': self.format_entry.extension if self.format_entry else Path(self.file_path).suffix.lower().lstrip('.'),
            },
            'format': self.format_entry.extension if self.format_entry else Path(self.file_path).suffix.lower().lstrip('.'),
        }

    def validate_security(self) -> None:
        from PIL import Image
        from PIL import UnidentifiedImageError

        try:
            with Image.open(self.file_path) as img:
                img.verify()
        except UnidentifiedImageError as exc:
            raise DocumentProcessingException(
                f'Image security validation failed: {exc}'
            ) from exc
        except Exception as exc:
            raise DocumentProcessingException(
                f'Image security validation failed: {exc}'
            ) from exc


# ------------------------------------------------------------------
# Spreadsheet adapter (XLSX)
# ------------------------------------------------------------------

class SpreadsheetAdapter(BaseDocumentAdapter):
    """
    Convert XLSX workbooks to a structured text representation.

    Preserves worksheet boundaries, headers, and row/cell data.
    Does not execute formulas or macros.
    """

    def normalize(self) -> dict[str, Any]:
        try:
            import openpyxl
        except ImportError as exc:
            raise UnsupportedFormatException(
                'openpyxl is not installed; XLSX processing is unavailable.'
            ) from exc

        wb = openpyxl.load_workbook(self.file_path, data_only=True, read_only=True)
        sheets = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            headers = []

            for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
                if row_idx == 1:
                    headers = [str(cell) if cell is not None else '' for cell in row]
                    continue
                row_data = {
                    str(headers[col_idx]): str(cell) if cell is not None else ''
                    for col_idx, cell in enumerate(row)
                    if col_idx < len(headers)
                }
                rows.append(row_data)

            sheets.append({
                'name': sheet_name,
                'headers': headers,
                'rows': rows,
                'row_count': len(rows),
            })

        wb.close()

        text_parts = []
        for sheet in sheets:
            text_parts.append(f'[Sheet: {sheet["name"]}]\n')
            if sheet['headers']:
                text_parts.append(' | '.join(sheet['headers']) + '\n')
            for row in sheet['rows']:
                text_parts.append(' | '.join(row.values()) + '\n')

        text = '\n'.join(text_parts)
        image_path = self._render_text_to_image(text)

        return {
            'text': text,
            'pages': [image_path] if image_path else [],
            'metadata': {
                'sheet_count': len(sheets),
                'sheets': [s['name'] for s in sheets],
                'format': 'xlsx',
            },
            'format': 'xlsx',
        }

    def _render_text_to_image(self, text: str) -> str | None:
        """
        Render extracted text to a PNG image for the vision pipeline.
        """
        try:
            from PIL import Image, ImageDraw, ImageFont

            font_size = 14
            margin = 20
            line_height = font_size + 4

            lines = text.split('\n')
            max_width = max(len(line) for line in lines) if lines else 80
            width = max(800, max_width * (font_size // 2) + margin * 2)
            height = max(600, len(lines) * line_height + margin * 2)

            img = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(img)

            try:
                font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', font_size)
            except Exception:
                font = ImageFont.load_default()

            y = margin
            for line in lines:
                draw.text((margin, y), line, fill='black', font=font)
                y += line_height

            output_dir = Path('/tmp/ocr_rendered')
            output_dir.mkdir(parents=True, exist_ok=True)
            image_path = output_dir / f'{self.upload_id}_xlsx.png'
            img.save(str(image_path))
            return str(image_path)
        except Exception as exc:
            logger.warning('Failed to render XLSX text to image: %s', exc)
            return None

    def validate_security(self) -> None:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(self.file_path, read_only=True)
            wb.close()
        except Exception as exc:
            raise DocumentProcessingException(
                f'XLSX security validation failed: {exc}'
            ) from exc


# ------------------------------------------------------------------
# CSV adapter
# ------------------------------------------------------------------

class CSVAdapter(BaseDocumentAdapter):
    """
    Convert CSV files to a structured text representation.
    """

    def normalize(self) -> dict[str, Any]:
        try:
            with open(self.file_path, 'r', encoding='utf-8', errors='replace') as f:
                sample = f.read(8192)
                f.seek(0)

                try:
                    dialect = csv.Sniffer().sniff(sample)
                except csv.Error:
                    dialect = csv.excel

                reader = csv.DictReader(f, dialect=dialect)
                headers = reader.fieldnames or []
                rows = []

                for row in reader:
                    rows.append({
                        str(k): str(v) if v is not None else ''
                        for k, v in row.items()
                    })

        except Exception as exc:
            raise DocumentProcessingException(
                f'Failed to parse CSV: {exc}'
            ) from exc

        text_parts = []
        if headers:
            text_parts.append(' | '.join(headers) + '\n')
        for row in rows:
            text_parts.append(' | '.join(row.values()) + '\n')

        text = '\n'.join(text_parts)
        image_path = self._render_text_to_image(text)

        return {
            'text': text,
            'pages': [image_path] if image_path else [],
            'metadata': {
                'column_count': len(headers),
                'row_count': len(rows),
                'headers': headers,
                'format': 'csv',
            },
            'format': 'csv',
        }

    def _render_text_to_image(self, text: str) -> str | None:
        """
        Render extracted text to a PNG image for the vision pipeline.
        """
        try:
            from PIL import Image, ImageDraw, ImageFont

            font_size = 14
            margin = 20
            line_height = font_size + 4

            lines = text.split('\n')
            max_width = max(len(line) for line in lines) if lines else 80
            width = max(800, max_width * (font_size // 2) + margin * 2)
            height = max(600, len(lines) * line_height + margin * 2)

            img = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(img)

            try:
                font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', font_size)
            except Exception:
                font = ImageFont.load_default()

            y = margin
            for line in lines:
                draw.text((margin, y), line, fill='black', font=font)
                y += line_height

            output_dir = Path('/tmp/ocr_rendered')
            output_dir.mkdir(parents=True, exist_ok=True)
            image_path = output_dir / f'{self.upload_id}_csv.png'
            img.save(str(image_path))
            return str(image_path)
        except Exception as exc:
            logger.warning('Failed to render CSV text to image: %s', exc)
            return None

    def validate_security(self) -> None:
        try:
            with open(self.file_path, 'rb') as f:
                header = f.read(16)
                if header.startswith(b'\x50\x4b\x03\x04'):
                    raise DocumentProcessingException(
                        'CSV file appears to be a ZIP archive.'
                    )
        except DocumentProcessingException:
            raise
        except Exception as exc:
            raise DocumentProcessingException(
                f'CSV security validation failed: {exc}'
            ) from exc


# ------------------------------------------------------------------
# Text adapter
# ------------------------------------------------------------------

class TextAdapter(BaseDocumentAdapter):
    """
    Read plain text files safely.

    Rejects binary content and enforces a reasonable size limit.
    """

    def normalize(self) -> dict[str, Any]:
        max_size = 5 * 1024 * 1024
        size = os.path.getsize(self.file_path)
        if size > max_size:
            raise DocumentProcessingException(
                f'Text file exceeds maximum size of {max_size} bytes.'
            )

        try:
            with open(self.file_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
        except Exception as exc:
            raise DocumentProcessingException(
                f'Failed to read text file: {exc}'
            ) from exc

        if '\x00' in text:
            raise DocumentProcessingException(
                'Text file contains null bytes — possible binary content.'
            )

        image_path = self._render_text_to_image(text)

        return {
            'text': text,
            'pages': [image_path] if image_path else [],
            'metadata': {
                'char_count': len(text),
                'line_count': text.count('\n') + 1,
                'format': 'txt',
            },
            'format': 'txt',
        }

    def _render_text_to_image(self, text: str) -> str | None:
        """
        Render extracted text to a PNG image for the vision pipeline.
        """
        try:
            from PIL import Image, ImageDraw, ImageFont

            font_size = 14
            margin = 20
            line_height = font_size + 4

            lines = text.split('\n')
            max_width = max(len(line) for line in lines) if lines else 80
            width = max(800, max_width * (font_size // 2) + margin * 2)
            height = max(600, len(lines) * line_height + margin * 2)

            img = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(img)

            try:
                font = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', font_size)
            except Exception:
                font = ImageFont.load_default()

            y = margin
            for line in lines:
                draw.text((margin, y), line, fill='black', font=font)
                y += line_height

            output_dir = Path('/tmp/ocr_rendered')
            output_dir.mkdir(parents=True, exist_ok=True)
            image_path = output_dir / f'{self.upload_id}_txt.png'
            img.save(str(image_path))
            return str(image_path)
        except Exception as exc:
            logger.warning('Failed to render text to image: %s', exc)
            return None

    def validate_security(self) -> None:
        max_size = 5 * 1024 * 1024
        size = os.path.getsize(self.file_path)
        if size > max_size:
            raise DocumentProcessingException(
                f'Text file exceeds maximum size of {max_size} bytes.'
            )


# ------------------------------------------------------------------
# Adapter factory
# ------------------------------------------------------------------

def get_adapter(file_path: str | Path, upload_id: str) -> BaseDocumentAdapter:
    """
    Return the correct adapter for the given file.

    Args:
        file_path: Path to the uploaded file.
        upload_id: UUID of the OCR upload record.

    Returns:
        A ``BaseDocumentAdapter`` subclass instance.

    Raises:
        UnsupportedFormatException: If the format is not supported.
    """
    from ocr.formats import detect_format

    fmt = detect_format(file_path)

    adapter_map = {
        'pdf': PDFAdapter,
        'docx': DOCXAdapter,
        'image': ImageAdapter,
        'spreadsheet': SpreadsheetAdapter,
        'csv': CSVAdapter,
        'text': TextAdapter,
    }

    adapter_cls = adapter_map.get(fmt.adapter)
    if adapter_cls is None:
        raise UnsupportedFormatException(
            f'No adapter registered for format "{fmt.adapter}".'
        )

    adapter = adapter_cls(file_path=file_path, upload_id=upload_id, format_entry=fmt)
    return adapter
